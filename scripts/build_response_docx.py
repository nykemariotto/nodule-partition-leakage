"""Transpose response_to_reviewers.md into the IEEE Access response-to-reviewers .docx template.

The template ships fixed slots (Reviewer#1 Concern #1-3, Reviewer#2 Concern #1-3). We have four
reviewers with 8/8/2/7 concerns, so the slots are regenerated rather than filled -- the STRUCTURE
the template asks for is preserved: (a) the reviewer's concern, (b) author response, (c) author
action.

Run with the project's own `nodules` env, which now carries python-docx. It previously said to use
`geral`, the ENMG project's environment: that worked but coupled this repository to an unrelated
project's dependencies, so a change there could break a deliverable here.
"""
import re
import sys

import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# RETIRED 2026-08-09. A coauthor revised the .docx directly -- flat edits, no tracked changes, plus
# 14 anchored comments -- so the .docx is now the source and this markdown is a stale derivative.
# Running this would overwrite that revision in silence, and the loss is not confined to the
# commented passages: the English was corrected throughout. The guard is a hard stop rather than a
# warning because the failure is invisible until someone reads the output and notices it reads like
# the old text.
sys.exit(
    "RETIRED: Response-to-Reviewers.docx is now the source, not response_to_reviewers.md.\n"
    "Running this script would overwrite a coauthor's revision. If the markdown pipeline is ever\n"
    "revived, first reconcile the .md against the .docx and delete this guard deliberately."
)

SRC = r"E:\NODULES\paper\2_resubmission\response_to_reviewers.md"
OUT = r"E:\NODULES\paper\2_resubmission\Response-to-Reviewers.docx"

ORIG_ID = "Access-2026-27906"
ORIG_TITLE = ("Patient-Level Versus Nodule-Level Data Partitioning in Pulmonary Nodule "
              "Classification: An Empirical Analysis of Performance Bias Across LNDb and "
              "LIDC-IDRI Datasets")

md = open(SRC, encoding="utf-8").read()

# ---------------------------------------------------------------- parse the markdown
def clean(t):
    t = re.sub(r"\s+", " ", t).strip()
    t = t.replace("**", "").replace("`", "")
    t = re.sub(r"\*(.+?)\*", r"\1", t)          # italics
    t = t.replace("§", "Section ")
    return t.strip()

def bullets_to_text(block):
    """Flatten a markdown bullet list into numbered sentences, preserving every sub-point."""
    out = []
    for b in re.findall(r"\n\s*- (.*?)(?=\n\s*- |\Z)", "\n" + block, re.S):
        t = clean(b).replace("→", ":")
        if t:
            out.append(t)
    return "  ".join(f"({i}) {t}" for i, t in enumerate(out, 1))


items = []          # (reviewer, n_within, concern, response, action)
# numbered items: **N.M —** (a) "..."  (b) ...  (c) ...
blocks = re.split(r"\n(?=\*\*\d+\.\d+ [—-])", md)
for b in blocks:
    # An item ENDS at the next top-level heading or horizontal rule. Without this the LAST item's
    # "(c)" -- terminated by \Z -- swallowed everything after it into the reviewer-facing body: the
    # change summary and the internal working section that follows it. The source file mixes both,
    # so this boundary is what keeps the generated document to the parts meant to be sent.
    b = re.split(r"\n(?=## |---\s*\n)", b)[0]
    # optional title: "**3.2 — Experiment 3.**" as well as plain "**3.1 —**"
    m = re.match(r"\*\*(\d+)\.(\d+) [—-]\s*([^*]*)\*\*", b)
    if not m:
        continue
    rev = int(m.group(1))
    body = b[m.end():]
    # Most items have separate (b) and (c). Item 3.2 -- the largest, and the one that matters most
    # since Reviewer 3 is the primary target -- writes "(b/c) Each sub-point:" followed by a bullet
    # list. The first version of this parser matched only "(b)" and "(c)" and silently dropped that
    # whole response, which is how it reached the .docx with an empty Author action.
    pa = re.search(r"\(a\)(.*?)(?=\n\(b)", body, re.S)
    combined = re.search(r"\n\(b/c\)(.*)", body, re.S)
    if combined:
        resp, act = "", bullets_to_text(combined.group(1))
    else:
        # (c) runs to the END of the item, not to the first blank line. It used to stop at "\n\n",
        # which silently truncated any action written as more than one paragraph -- item 4.4 lost
        # every number it reported that way. The blocks were already split on the "**N.M --**"
        # marker before we get here, so \Z is the correct terminator.
        pb = re.search(r"\n\(b\)(.*?)(?=\n\(c\)|\Z)", body, re.S)
        pc = re.search(r"\n\(c\)(.*)", body, re.S)
        resp = clean(pb.group(1)) if pb else ""
        act = clean(pc.group(1)) if pc else ""
    items.append((rev, int(m.group(2)), clean(pa.group(1)) if pa else "", resp, act))

# Reviewer 2 used to be a bullet list and was special-cased here. Since
# 2026-08-04 it is numbered (a)/(b)/(c) items like every other reviewer, so the
# generic parser above already covers it; keeping the special case would have
# emitted its ten concerns twice.


# ---------------------------------------------------------------- build the document
d = docx.Document()
st = d.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(11)


def para(text="", bold=False, italic=False, space_after=6, colour=None):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    if colour:
        r.font.color.rgb = colour
    p.paragraph_format.space_after = Pt(space_after)
    return p


RED = RGBColor(0xC0, 0x00, 0x00)

para(f"Original Manuscript ID: {ORIG_ID}", bold=True, space_after=2)
para(f"Original Article Title: \u201c{ORIG_TITLE}\u201d", bold=True, space_after=2)
para("New Article Title: “Slice-Level Data Partitioning Inflates Pulmonary Nodule "
     "Classification Performance: A Controlled Three-Arm Study”", bold=True, space_after=10)

para("To: IEEE Access Editor", space_after=2)
para("Re: Response to reviewers", space_after=10)

para("Dear Editor,", space_after=6)

para("Thank you for allowing a resubmission of our manuscript, with an opportunity to address the "
     "reviewers' comments.", space_after=6)

# --- the title-change declaration (D63) -------------------------------------------------
para("We note at the outset that the title of this resubmission differs from the original. The "
     "change follows directly from Reviewer 3's central criticism, which observed that \u201cthe "
     "title, abstract, discussion, and conclusion attribute a reported 6\u201314 percentage-point "
     "difference to patient-level versus nodule-level partitioning\u201d while \u201cno matched "
     "nodule-level experiment is performed using the same cohort, labels, preprocessing, "
     "architecture, training procedure, and evaluation set.\u201d We have now performed that "
     "matched experiment. It shows that the contrast named in the original title is not where the "
     "inflation arises, so retaining that title would repeat the very mismatch between claim and "
     "evidence that Reviewer 3 identified. The study, the cohort, the research question and the "
     "author list are unchanged.", space_after=8)

# --- opening letter, carried over from the markdown -------------------------------------
op = re.search(r"## Opening letter.*?\n\n(.*?)\n\n---", md, re.S)
if op:
    for chunk in op.group(1).split("\n\n"):
        para(clean(chunk), space_after=6)

para("We are uploading (a) our point-by-point response to the comments (below), (b) an updated "
     "manuscript with the changes highlighted, and (c) a clean updated manuscript.", space_after=6)
para("Best regards,", space_after=2)
para("N. Mariotto et al.", space_after=12)

para("\u2014" * 30, space_after=10)

cur = None
for rev, n, concern, resp, action in items:
    if rev != cur:
        cur = rev
        h = d.add_paragraph()
        hr = h.add_run(f"Reviewer #{rev}")
        hr.bold = True
        hr.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)

    para(f"Reviewer #{rev}, Concern # {n}:", bold=True, space_after=2)
    para(concern, italic=True, space_after=4)
    if resp:
        p = d.add_paragraph()
        p.add_run("Author response: ").bold = True
        p.add_run(resp)
        p.paragraph_format.space_after = Pt(4)
    if action:
        p = d.add_paragraph()
        p.add_run("Author action: ").bold = True
        p.add_run(action)
        p.paragraph_format.space_after = Pt(10)

d.add_page_break()
para("NOTE FOR THE AUTHORS \u2014 DELETE BEFORE SUBMITTING", bold=True, colour=RED, space_after=6)
para("This file was generated from response_to_reviewers.md. Everything above is factual and "
     "traceable to artifacts in the repository. Before submitting:", space_after=6)
for t in [
    "1. Revise the CRediT statement with the coauthors \u2014 it still describes the original study.",
    "2. Decide the bibliography packaging: the .bib travelling with the source, or the generated "
    ".bbl pasted inline (the original submission used hand-written \\bibitem entries).",
    "3. Re-read for grammar, and confirm references print in order of citation.",
    "4. Delete this page.",
]:
    para(t, space_after=4)

d.save(OUT)
print("wrote", OUT)
print("concerns transposed:", len(items))
